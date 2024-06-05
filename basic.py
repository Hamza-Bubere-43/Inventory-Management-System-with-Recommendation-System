
import os
import numpy as np
from collections import defaultdict
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import re


def extract_data_from_file(file_path):
    # Initialize variables to store extracted data
    date = None
    time = None
    items = []
    total_cost = None
    physical_address = None
    distance = None
    
    # Open the file and read its content
    with open(file_path, 'r') as file:
        lines = file.readlines()
    
    # Extract data from each line
    for line_index, line in enumerate(lines):
        if "Date:" in line:
            # Extract date and time using regular expressions
            date_match = re.search(r'Date: (\d{4}-\d{2}-\d{2}) (\d{2}:\d{2}:\d{2})', line)
            if date_match:
                date = date_match.group(1)
                time = date_match.group(2)
        elif "Items:" in line:
            # Extract item details
            for next_line in lines[line_index + 2:]:
                if next_line.strip() == "":
                    break  # End of item section
                item_match = re.match(r'(\w+(?:\s\w+)*)\s+(\d+\.\d+)\s+(\$[\d.]+)', next_line.strip())
                if item_match:
                    item_name = item_match.group(1)
                    amount = float(item_match.group(2))
                    cost = float(item_match.group(3).replace('$', ''))
                    items.append((item_name, amount, cost))
        elif "Total Cost:" in line:
            # Extract total cost
            total_cost_match = re.search(r'\$([\d.]+)', lines[line_index + 1])
            if total_cost_match:
                total_cost = float(total_cost_match.group(1))
        elif "Physical Address:" in line:
            # Extract physical address
            physical_address = line.split(":")[1].strip()
        elif "distance:" in line.lower():  # Convert line to lower case for case-insensitive matching
            # Extract distance
            distance_match = re.search(r'distance:\s*([\d.]+)\s*km', line, re.IGNORECASE)
            if distance_match:
                distance = float(distance_match.group(1))

    return {
        "Date": date,
        "Time": time,
        "Items": items,
        "Total Cost": total_cost,
        "Physical Address": physical_address,
        "Distance": distance
    }

def extract_data_from_directory(directory_path):
    # Initialize dictionary to store extracted data for each file
    all_data = {}

    # Iterate over each file in the directory
    for file_name in os.listdir(directory_path):
        if file_name.endswith(".txt"):
            # Construct full file path
            file_path = os.path.join(directory_path, file_name)
            
            # Extract data from the file
            data = extract_data_from_file(file_path)
            
            # Store the extracted data with the file name as the key
            all_data[file_name] = data
    
    # Sort the dictionary by keys (dates)
    sorted_all_data = dict(sorted(all_data.items(), key=lambda x: x[1]["Date"]))
    
    # Extract sorted dates, total costs, and distances
    dates = [data["Date"] for data in sorted_all_data.values()]
    total_costs = [data["Total Cost"] for data in sorted_all_data.values()]
    distances = [data["Distance"] for data in sorted_all_data.values()]

    return sorted_all_data, dates, total_costs, distances


def ans():
    # Get the current working directory
    cwd = os.getcwd()

    # Update the directory path to point to the 'bills' folder in the current working directory
    directory_path = os.path.join(cwd, "bills")
    sorted_all_data, dates, total_costs, distances = extract_data_from_directory(directory_path)

    # Plotting distances over dates
    plt.plot(dates, distances)
    plt.xlabel("Date")
    plt.ylabel("Distance (km)")
    plt.title("Distance Travelled Over Time")
    plt.xticks(rotation=45)
    plt.tight_layout()
    return plt

# Function to calculate total money made per day
def calculate_money_per_day_from_data(sorted_all_data):
    money_per_day = {}
    
    for data in sorted_all_data.values():
        date = data["Date"]
        total_cost = data["Total Cost"]
        
        if date in money_per_day:
            money_per_day[date] += total_cost if total_cost else 0
        else:
            money_per_day[date] = total_cost if total_cost else 0
    
    return money_per_day

# Function to plot money made per day
def plot_money_per_day_from_data(money_per_day):
    dates = list(money_per_day.keys())
    total_money = list(money_per_day.values())

    plt.figure(figsize=(10, 6))
    plt.plot(dates, total_money, marker='o', linestyle='-')
    plt.title('Total Money Made per Day')
    plt.xlabel('Date')
    plt.ylabel('Total Money ($)')
    plt.xticks(rotation=45)
    plt.tight_layout()
    return plt

# Main function
def process_bills_data():
    import os

    # Get the current working directory
    cwd = os.getcwd()

    # Update the directory path to point to the 'bills' folder in the current working directory
    directory_path = os.path.join(cwd, "bills")
    print("Extracting data from text files...")
    
    # Extract data from all text files in the directory
    sorted_all_data, dates, total_costs, _ = extract_data_from_directory(directory_path)
    print("Data extraction completed.")
    
    # Print extracted dates and total costs
    print("\nExtracted Dates and Total Costs:")
    # for date, cost in zip(dates, total_costs):
    #     print(f"Date: {date}, Total Cost: {cost}")
    
    # Calculate total money made per day
    print("\nCalculating total money made per day...")
    money_per_day = calculate_money_per_day_from_data(sorted_all_data)
    print("Calculation completed.")
    
    # Plot total money made per day
    print("\nPlotting total money made per day...")
    plt = plot_money_per_day_from_data(money_per_day)
    print("Plotting completed.")
    return plt


def extract_information():
    cwd = os.getcwd()
    directory = os.path.join(cwd, "destocklog")
    date_ingredient_weights = defaultdict(lambda: defaultdict(float))
    plots = []

    for filename in os.listdir(directory):
        if filename.endswith('.txt'):
            file_path = os.path.join(directory, filename)
            date, ingredients = extract_from_file12(file_path)
            if date:
                for ingredient, weight in ingredients.items():
                    date_ingredient_weights[date][ingredient] += weight

    date_ingredient_weights = {
        datetime.strptime(date, '%Y-%m-%d'): ingredients
        for date, ingredients in date_ingredient_weights.items() if date
    }

    max_date = max(date_ingredient_weights.keys())

    delta = timedelta(days=7)

    # Store only the latest 5 plots
    latest_plots = []

    for i in range(5):
        end_date = max_date - delta * i
        start_date = end_date - timedelta(days=6)

        ingredients_for_period = defaultdict(float)
        for date in date_ingredient_weights.keys():
            if start_date <= date <= end_date:
                for ingredient, weight in date_ingredient_weights[date].items():
                    ingredients_for_period[ingredient] += weight

        top_20_ingredients = sorted(
            ingredients_for_period.items(), key=lambda x: x[1], reverse=True)[:20]

        ingredient_names = [ingredient[0] for ingredient in top_20_ingredients]
        ingredient_weights = [ingredient[1] for ingredient in top_20_ingredients]

        plt.figure(figsize=(12, 6))
        plt.barh(ingredient_names, ingredient_weights, color='skyblue')
        plt.xlabel('Weight (kg)')
        plt.ylabel('Ingredient')
        plt.title(
            f'Top 20 Ingredients from {start_date.strftime("%d-%m-%Y")} to {end_date.strftime("%d-%m-%Y")}')
        plt.xticks(rotation=45)
        plt.tight_layout()

        latest_plots.append(plt)

    return latest_plots

def extract_from_file12(file_path):
    date = None
    ingredient_dict = {}

    with open(file_path, 'r') as file:
        data = file.read()

    date_match = re.search(
        r'Date: (\d{4}-\d{2}-\d{2}) \d{2}:\d{2}:\d{2}', data)
    if date_match:
        date = date_match.group(1)

    ingredients = re.findall(r'([A-Za-z\s]+):\s([\d\.]+)\s[kKgG]+', data)

    for ingredient, weight in ingredients:
        ingredient_dict[ingredient.strip()] = float(weight)

    return date, ingredient_dict

def get_date_ranges(start_date, end_date, delta):
    curr_date = start_date
    while curr_date < end_date:
        yield curr_date, min(curr_date + delta, end_date)
        curr_date += delta


import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def gen_report():
    # Create a folder called 'report2' if it doesn't exist
    folder_name = 'report2'
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)

    # Create a new Word document
    doc = Document()

    # Assuming m1, m2, m3 are Matplotlib figures or arrays of Matplotlib figures
    m1 = ans()
    m2 = process_bills_data()
    m3 = extract_information()

    q = []
    q.append(m1)
    q.append(m2)
    for i in m3:
        q.append(i)
        
        
    # for i in q:
    #     i.show()
    # Iterate over the plots in the list
    for index, fig in enumerate(q):
        # Save the plot as an image
        fig_path = os.path.join(folder_name, f"plot_{index}.png")
        fig.savefig(fig_path, bbox_inches='tight', dpi=300)

        # Add a new page to the document
        doc.add_page_break()

        # Add the image to the document
        doc.add_heading(f"Plot {index+1}", 0)
        doc.add_picture(fig_path, width=Inches(6.5))  # Adjust the width as needed

        # Close the current figure
        plt.close()

    # Save the Word document
    file_name = os.path.join(folder_name, "report.docx")
    index = 1
    while os.path.exists(file_name):
        file_name = os.path.join(folder_name, f"report_{index}.docx")
        index += 1
    doc.save(file_name)




#gen_report()